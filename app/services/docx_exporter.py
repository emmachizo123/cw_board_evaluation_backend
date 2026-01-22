"""
app.services.docx_exporter
-------------------------
Generate a board-ready DOCX report from a saved report JSON payload.

Expected input shape (saved report payload):
{
  "report_id": "...",
  "evaluation_id": "...",
  "status": "draft",
  "summary_json": {
     "analytics": {...},
     "context": {...},
     "sections": {...}
  }
}
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


OUTPUT_DOCX_DIR = Path("outputs/docx")
OUTPUT_DOCX_DIR.mkdir(parents=True, exist_ok=True)

PRIORITY_LABELS = {
    "quick_win": "Quick Win (0–90 days)",
    "near_term": "Near Term (3–6 months)",
    "structural": "Structural (6–12 months)",
    "long_term": "Long Term (12+ months)",
}


def export_report_to_docx(report_payload: Dict[str, Any]) -> Path:
    """
    Create a DOCX file from a report payload.

    Parameters
    ----------
    report_payload : dict
        The saved report object returned by load_report_by_id() or load_latest_report_for_evaluation().

    Returns
    -------
    Path
        Path to the generated DOCX file.
    """
    report_id = report_payload.get("report_id", "unknown")
    evaluation_id = report_payload.get("evaluation_id", "unknown")
    status = report_payload.get("status", "draft")
    summary = report_payload.get("summary_json") or {}

    analytics = summary.get("analytics") or {}
    context = summary.get("context") or {}
    sections = summary.get("sections") or {}

    exec_summary = sections.get("executive_summary") or {}
    recs_block = sections.get("recommendations") or {}
    recs = (recs_block.get("recommendations") or [])

    doc = Document()
    _set_default_styles(doc)

    # ---- Cover / Title ----
    title = doc.add_paragraph("Board Evaluation Report")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(20)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ---- Meta ----
    meta = doc.add_paragraph()
    meta.add_run("Evaluation ID: ").bold = True
    meta.add_run(str(evaluation_id))
    meta.add_run("\nReport ID: ").bold = True
    meta.add_run(str(report_id))
    meta.add_run("\nStatus: ").bold = True
    meta.add_run(str(status))
    meta.add_run("\nGenerated: ").bold = True
    meta.add_run(datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC"))

    doc.add_page_break()

    # ---- Executive Summary ----
    doc.add_heading("Executive Summary", level=1)
    _para(doc, exec_summary.get("overall_message"))

    strengths = exec_summary.get("key_strengths") or []
    if strengths:
        doc.add_heading("Key Strengths", level=2)
        _bullets(doc, strengths)

    weaknesses = exec_summary.get("key_weaknesses") or []
    if weaknesses:
        doc.add_heading("Key Weaknesses", level=2)
        _bullets(doc, weaknesses)

    outlook = exec_summary.get("outlook")
    if outlook:
        doc.add_heading("Outlook", level=2)
        _para(doc, outlook)

    # ---- Evaluation Snapshot ----
    doc.add_page_break()
    doc.add_heading("Evaluation Snapshot", level=1)
    _snapshot_table(doc, analytics)

    # ---- Trends ----
    trends = (analytics.get("trends") or {})
    if trends.get("available"):
        doc.add_paragraph()
        doc.add_heading("Historical Trend (Overall Score)", level=2)
        _trend_table(doc, trends)

    # ---- Recommendations ----
    doc.add_page_break()
    doc.add_heading("Recommendations", level=1)
    if recs:
        doc.add_paragraph(
            "The following actions are prioritised into time horizons and mapped to governance outcomes."
        )
        _recommendations_table(doc, recs)
    else:
        _para(doc, "No recommendations were generated.")

    # ---- Appendix: References used ----
    doc.add_page_break()
    doc.add_heading("Appendix: References Used", level=1)
    _references_section(doc, context)

    # Save
    import tempfile
    from uuid import uuid4

    # Save (unique name to avoid Windows/Word lock issues)
    ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    out_path = OUTPUT_DOCX_DIR / f"report__{evaluation_id}__{report_id}__{ts}.docx"
    doc.save(str(out_path))
    return out_path


# ----------------- formatting helpers -----------------

def _set_default_styles(doc: Document) -> None:
    """Set default font/size."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _para(doc: Document, text: Any) -> None:
    """Add a paragraph if text is present."""
    t = (text or "").strip()
    doc.add_paragraph(t if t else "")


def _bullets(doc: Document, items: List[Any]) -> None:
    """Bullet list."""
    for it in items:
        t = str(it).strip()
        if t:
            doc.add_paragraph(t, style="List Bullet")


def _snapshot_table(doc: Document, analytics: Dict[str, Any]) -> None:
    """Snapshot summary + dimension scores table."""
    ev = analytics.get("evaluation") or {}
    metrics = analytics.get("metrics") or {}
    dims = metrics.get("dimensions") or []
    stats = analytics.get("response_stats") or {}

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    def row(k: str, v: Any):
        r = table.add_row().cells
        r[0].text = k
        r[1].text = "" if v is None else str(v)

    row("Client", ev.get("tenant_name"))
    row("Sector", ev.get("sector"))
    row("Year", ev.get("year"))
    row("Regulators", ", ".join(ev.get("regulators") or []))
    row("Invited", stats.get("invited"))
    row("Responded", stats.get("responded"))
    row("Completion Rate", stats.get("completion_rate"))
    row("Overall Score", metrics.get("overall_score"))

    doc.add_paragraph()
    doc.add_heading("Dimension Scores", level=2)

    dim_table = doc.add_table(rows=1, cols=2)
    dim_table.style = "Table Grid"
    dim_table.rows[0].cells[0].text = "Dimension"
    dim_table.rows[0].cells[1].text = "Score"

    for d in dims:
        r = dim_table.add_row().cells
        r[0].text = str(d.get("name", ""))
        r[1].text = str(d.get("score", ""))


def _trend_table(doc: Document, trends: Dict[str, Any]) -> None:
    """Simple trend table (year vs overall score)."""
    years = trends.get("years") or []
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.rows[0].cells[0].text = "Year"
    t.rows[0].cells[1].text = "Overall Score"

    for y in years:
        r = t.add_row().cells
        r[0].text = str(y.get("year", ""))
        r[1].text = str(y.get("overall_score", ""))


def _recommendations_table(doc: Document, recs: List[Dict[str, Any]]) -> None:
    """Board-friendly recommendations table."""
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text = "Theme"
    hdr[1].text = "Priority"
    hdr[2].text = "Action"
    hdr[3].text = "Owner"
    hdr[4].text = "Timeline"
    hdr[5].text = "Success Metric"

    for r in recs:
        row = table.add_row().cells
        row[0].text = str(r.get("theme", ""))
        row[1].text = PRIORITY_LABELS.get(str(r.get("priority", "")).strip(), str(r.get("priority", "")))
        row[2].text = str(r.get("action", ""))
        row[3].text = str(r.get("owner_suggestion", ""))
        row[4].text = str(r.get("timeline", ""))
        row[5].text = str(r.get("success_metric", ""))


def _references_section(doc: Document, context: Dict[str, Any]) -> None:
    """Appendix section: regulatory snippets + C&W playbook snippets."""
    regs = context.get("regulatory_snippets") or []
    playbook = context.get("cw_playbook_snippets") or []

    if regs:
        doc.add_heading("Regulatory References", level=2)
        for r in regs:
            label = f"{r.get('framework_code', '')} ({r.get('theme', '')})"
            doc.add_paragraph(label, style="List Bullet")
            txt = (r.get("text") or "").strip()
            if txt:
                doc.add_paragraph(txt)

    if playbook:
        doc.add_heading("C&W Methodology / Playbook References", level=2)
        for p in playbook:
            src = (p.get("source") or "").strip()
            txt = (p.get("text") or "").strip()
            if src:
                doc.add_paragraph(src, style="List Bullet")
            if txt:
                doc.add_paragraph(txt)
